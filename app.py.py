from flask import Flask, render_template, request, send_file, jsonify
from datetime import datetime, timedelta
from docx import Document
import io
import csv

app = Flask(__name__)

# Define default deadlines
default_deadlines = {
    "Join additional parties": 60,
    "Amend pleadings": 7,
    "File motions to limit/suppress/exclude evidence or testimony": 30,
    "Last day to propound discovery to opposing party": 60,
    "Complete all other discovery": 30,
    "Conduct mediation": 30,
    "Complete depositions of fact witnesses": 30,
    "Complete depositions of experts": 30,
    "Designate expert witnesses for affirmative relief": 120,
    "Designate other expert witnesses": 90,
    "Designate rebuttal expert witnesses": 15,
    "Submit proposed findings of fact and conclusions of law": 10,
    "Submit proposed jury charges, instructions and definitions": 7,
    "File business records affidavit under TRE 902(10)": 14,
}

@app.route('/')
def index():
    return render_template('index.html', deadlines=default_deadlines)

@app.route('/calculate', methods=['POST'])
def calculate():
    trial_date_str = request.form['trial_date']
    trial_date = datetime.strptime(trial_date_str, '%Y-%m-%d')

    calculated = []
    for event, days in default_deadlines.items():
        deadline_date = trial_date - timedelta(days=days)
        calculated.append({"event": event, "date": deadline_date.strftime('%Y-%m-%d')})

    return jsonify(calculated)

@app.route('/generate', methods=['POST'])
def generate():
    trial_date_str = request.form['trial_date']
    client_name = request.form['client_name']
    trial_date = datetime.strptime(trial_date_str, '%Y-%m-%d')

    # Get the edited deadlines from the form
    events = request.form.getlist('event')
    dates = request.form.getlist('date')

    # Create Word doc
    doc = Document()
    doc.add_heading('Texas Family Law Trial Deadlines', 0)
    doc.add_paragraph(f'Client: {client_name}')
    doc.add_paragraph(f'Trial Date: {trial_date.strftime("%B %d, %Y")}')
    doc.add_paragraph('')

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Event'
    hdr_cells[1].text = 'Deadline'

    for event, date in zip(events, dates):
        row_cells = table.add_row().cells
        row_cells[0].text = event
        row_cells[1].text = datetime.strptime(date, '%Y-%m-%d').strftime('%B %d, %Y')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name=f'{client_name}_Deadlines.docx')
from flask import send_file
from datetime import datetime
import io

@app.route('/download_ics', methods=['POST'])
def download_ics():
    client_name = request.form.get('client_name', '').strip()
    events = request.form.getlist('event')
    dates = request.form.getlist('date')

    def esc(s: str) -> str:
        # iCalendar escaping of \ ; , characters
        return s.replace('\\', '\\\\').replace(';', r'\;').replace(',', r'\,')

    lines = [
        'BEGIN:VCALENDAR',
        'VERSION:2.0',
        'CALSCALE:GREGORIAN',
        'PRODID:-//Freelance Annie//Deadline Tool//EN',
    ]

    now_utc = datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')

    for event, date_str in zip(events, dates):
        # Expecting YYYY-MM-DD coming from your <input type="date">
        ymd = datetime.strptime(date_str, '%Y-%m-%d').strftime('%Y%m%d')
        summary = f'{client_name} – {event}' if client_name else event
        uid = f'{ymd}-{abs(hash(event))}@freelanceannie'
        lines += [
            'BEGIN:VEVENT',
            f'UID:{uid}',
            f'DTSTAMP:{now_utc}',
            f'SUMMARY:{esc(summary)}',
            f'DTSTART;VALUE=DATE:{ymd}',
            f'DTEND;VALUE=DATE:{ymd}',  # all-day event
            'END:VEVENT',
        ]

    lines.append('END:VCALENDAR')

    buf = io.BytesIO('\r\n'.join(lines).encode('utf-8'))
    filename = f"{client_name or 'deadlines'}.ics"
    return send_file(buf, as_attachment=True, download_name=filename, mimetype='text/calendar')

@app.route('/download_csv', methods=['POST'])
def download_csv():
    client_name = request.form['client_name']
    events = request.form.getlist('event')
    dates = request.form.getlist('date')

    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(['Subject', 'Start Date', 'Start Time', 'End Date', 'End Time', 'Description'])

    for event, date in zip(events, dates):
        writer.writerow([event, datetime.strptime(date, '%Y-%m-%d').strftime('%m/%d/%Y'), "09:00 AM", datetime.strptime(date, '%Y-%m-%d').strftime('%m/%d/%Y'), "09:30 AM", f'{event} Deadline'])

    buffer.seek(0)
    return send_file(io.BytesIO(buffer.getvalue().encode()), as_attachment=True, download_name=f'{client_name}_Deadlines.csv', mimetype='text/csv')

if __name__ == '__main__':
    app.run(debug=True)

