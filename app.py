from flask import Flask, render_template, request, send_file, jsonify
from datetime import datetime, timedelta
from docx import Document
import io
import csv

app = Flask(__name__)

# Default trial-related deadlines (days BEFORE trial)
default_deadlines = {
    "Join additional parties": 60,
    "Amend pleadings": 7,
    "File motions to limit/suppress/exclude evidence or testimony": 30,
    "Last day to propound discovery to opposing party": 60,
    "Complete all other discovery": 30,
    "Conduct mediation": 30,
    "Complete depositions of fact witnesses": 30,
    "Complete depositions of experts": 30,
    "Designate expert witnesses for affirmative relief": 120,
    "Designate other expert witnesses": 90,
    "Designate rebuttal expert witnesses": 15,
    "Submit proposed findings of fact and conclusions of law": 10,
    "Submit proposed jury charges, instructions and definitions": 7,
    "File business records affidavit under TRE 902(10)": 14,
}


@app.route('/')
def index():
    """Render main calculator page."""
    return render_template('index.html')


@app.route('/calculate', methods=['POST'])
def calculate():
    """Return JSON list of deadlines based on the given trial date."""
    trial_date_str = request.form.get('trial_date')
    if not trial_date_str:
        return jsonify([])

    try:
        trial_date = datetime.strptime(trial_date_str, '%Y-%m-%d').date()
    except ValueError:
        return jsonify([])

    results = []
    for event, days_before in default_deadlines.items():
        deadline_date = trial_date - timedelta(days=days_before)
        results.append({
            'event': event,
            'date': deadline_date.strftime('%Y-%m-%d'),
        })
    return jsonify(results)


@app.route('/download_ics', methods=['POST'])
def download_ics():
    """Generate an .ics calendar file from posted events and dates."""
    events = request.form.getlist('event')
    dates = request.form.getlist('date')
    client_name = request.form.get('client_name', '').strip()

    def esc(value: str) -> str:
        # Escape characters per RFC 5545
        return (value or '').replace('\\', '\\\\').replace(';', '\\;').replace(',', '\\,').replace('\n', '\\n')

    lines = [
        'BEGIN:VCALENDAR',
        'VERSION:2.0',
        'PRODID:-//Freelance Annie//Texas Deadline Calculator//EN',
    ]

    now_utc = datetime.utcnow().strftime('%Y%m%dT%H%M%SZ')

    for event, date_str in zip(events, dates):
        if not date_str:
            continue
        try:
            ymd = datetime.strptime(date_str, '%Y-%m-%d').strftime('%Y%m%d')
        except ValueError:
            continue

        summary = f'{client_name} – {event}' if client_name else event
        uid = f'{ymd}-{abs(hash((event, date_str)))}@freelanceannie'

        lines.extend([
            'BEGIN:VEVENT',
            f'UID:{uid}',
            f'DTSTAMP:{now_utc}',
            f'SUMMARY:{esc(summary)}',
            f'DTSTART;VALUE=DATE:{ymd}',
            f'DTEND;VALUE=DATE:{ymd}',
            'END:VEVENT',
        ])

    lines.append('END:VCALENDAR')
    ics_data = '\r\n'.join(lines) + '\r\n'

    buf = io.BytesIO(ics_data.encode('utf-8'))
    filename = f"{client_name or 'Deadlines'}-deadlines.ics"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype='text/calendar',
    )


@app.route('/download_csv', methods=['POST'])
def download_csv():
    """Generate a CSV suitable for calendar import (e.g., Google Calendar)."""
    events = request.form.getlist('event')
    dates = request.form.getlist('date')
    client_name = request.form.get('client_name', '').strip()

    buffer = io.StringIO()
    writer = csv.writer(buffer)

    # Google Calendar style: Subject,Start Date,Start Time,End Date,End Time,Description
    writer.writerow(['Subject', 'Start Date', 'Start Time', 'End Date', 'End Time', 'Description'])

    for event, date_str in zip(events, dates):
        if not date_str:
            continue
        try:
            dt = datetime.strptime(date_str, '%Y-%m-%d')
        except ValueError:
            continue

        date_mdy = dt.strftime('%m/%d/%Y')
        subject = f'{client_name} – {event}' if client_name else event
        description = f'{event} Deadline'

        writer.writerow([subject, date_mdy, '09:30 AM', date_mdy, '09:30 AM', description])

    buffer.seek(0)
    filename = f"{client_name or 'Deadlines'}-deadlines.csv"
    return send_file(
        io.BytesIO(buffer.getvalue().encode('utf-8')),
        as_attachment=True,
        download_name=filename,
        mimetype='text/csv',
    )


@app.route('/generate', methods=['POST'])
def generate_word():
    """Optional server-side Word document generator (not required for basic use)."""
    events = request.form.getlist('event')
    dates = request.form.getlist('date')
    client_name = request.form.get('client_name', '').strip()
    trial_date = request.form.get('trial_date', '').strip()

    doc = Document()
    doc.add_heading('Legal Deadlines Summary', level=1)

    if client_name:
        doc.add_paragraph(f'Client: {client_name}')
    if trial_date:
        doc.add_paragraph(f'Trial Date: {trial_date}')

    table = doc.add_table(rows=1, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Event'
    hdr_cells[1].text = 'Deadline'

    for event, date in zip(events, dates):
        if not event and not date:
            continue
        row_cells = table.add_row().cells
        row_cells[0].text = event
        row_cells[1].text = date

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{client_name or 'Deadlines'}-deadlines.docx"
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


if __name__ == '__main__':
    app.run(debug=True)
